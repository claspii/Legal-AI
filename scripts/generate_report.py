import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt lề trong (padding) cho ô trong bảng (đơn vị dxa, 1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_callout(doc, text):
    """Tạo một đoạn callout nổi bật giống dạng box trong các báo cáo chuyên nghiệp."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    # Tạo viền trái bằng XML chèn trực tiếp vào paragraph properties (Màu đen)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="12" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # Nền xám nhẹ cho callout
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F8"/>')
    pPr.append(shd)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14.5)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

def main():
    doc = docx.Document()
    
    # Thiết lập lề trang chuẩn (1 inch = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Màu sắc chủ đạo - Tất cả màu đen theo yêu cầu
    PRIMARY_COLOR = RGBColor(0, 0, 0)
    TEXT_COLOR = RGBColor(0, 0, 0)
    ACCENT_COLOR = RGBColor(0, 0, 0)
    
    # Cấu hình Default Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(15)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 11)
    style_normal.font.color.rgb = TEXT_COLOR
    
    # -------------------------------------------------------------
    # TIÊU ĐỀ CHÍNH (Căn giữa, Chữ in hoa, Màu đen)
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(6)
    
    title_run = title_p.add_run("BÁO CÁO TIẾN ĐỘ ĐỒ ÁN")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(26)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    # Phụ đề
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle_p.add_run("Phát triển Hệ thống Legal RAG + Đồ thị Tri thức Neo4j và Chưng cất Dữ liệu")
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(17)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = ACCENT_COLOR
    
    # -------------------------------------------------------------
    # 1. NỘI DUNG ĐÃ HOÀN THÀNH
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True
    run_h1 = h1.add_run("1. Nội dung đã hoàn thành")
    run_h1.font.name = 'Times New Roman'
    run_h1.font.size = Pt(20)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = PRIMARY_COLOR
    
    # -------------------------------------------------------------
    # 1.1. Xây dựng Cơ sở dữ liệu đồ thị (Graph Database)
    # -------------------------------------------------------------
    h1_1 = doc.add_paragraph()
    h1_1.paragraph_format.space_before = Pt(12)
    h1_1.paragraph_format.space_after = Pt(4)
    h1_1.paragraph_format.keep_with_next = True
    run_h1_1 = h1_1.add_run("1.1. Xây dựng Cơ sở dữ liệu đồ thị (Graph Database)")
    run_h1_1.font.name = 'Times New Roman'
    run_h1_1.font.size = Pt(17)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 13)
    run_h1_1.font.bold = True
    run_h1_1.font.color.rgb = ACCENT_COLOR
    
    p = doc.add_paragraph(
        "Nhóm phát triển đã hoàn thành việc thiết lập và xây dựng Cơ sở dữ liệu đồ thị (Knowledge Graph) "
        "trên nền tảng Neo4j kết hợp cơ sở dữ liệu Vector trên ChromaDB. Dữ liệu luật được xử lý từ 4 bộ luật "
        "gốc tiếng Việt dưới định dạng văn bản thô bao gồm: Dân sự, Hình sự, Hôn nhân & Gia đình và Lao động. "
        "Quy trình xây dựng đồ thị sử dụng cách tiếp cận Hybrid:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    bullets = [
        ("Trích xuất cấu trúc (Rule-based): ", "Sử dụng các biểu thức chính quy (Regex) tối ưu để bóc tách cấu trúc phân cấp chặt chẽ của luật Việt Nam bao gồm các cấp độ Luật, Chương, Điều, Khoản và các mối liên kết cấu trúc (HAS_CHAPTER, HAS_ARTICLE, HAS_CLAUSE) cũng như các tham chiếu chéo giữa các điều luật (REFERENCES, CROSS_LAW_REF)."),
        ("Trích xuất ngữ nghĩa (LLM-based): ", "Sử dụng mô hình ngôn ngữ lớn (LLM) để phân tích nội dung từng chunk tài liệu nhằm bóc tách các thực thể ngữ nghĩa bao gồm Khái niệm pháp lý (Concept), Chủ thể (Actor), Hành vi pháp lý (Action) và các mối quan hệ ngữ nghĩa liên quan (DEFINES, MENTIONS, RELATED_TO, REGULATES, PERFORMS) để làm giàu tri thức của hệ thống.")
    ]
    for bold_text, normal_text in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        
        r_bold = bp.add_run(bold_text)
        r_bold.font.name = 'Times New Roman'
        r_bold.bold = True
        r_bold.font.color.rgb = TEXT_COLOR
        
        r_norm = bp.add_run(normal_text)
        r_norm.font.name = 'Times New Roman'
        r_norm.font.color.rgb = TEXT_COLOR
        
    add_callout(
        doc, 
        "Hệ thống đã thiết lập các ràng buộc duy nhất (Unique Constraints) cho các thuộc tính khóa của các Node thuộc nhãn "
        "Law, Chapter, Article, Clause, Concept, Actor, Action và Chunk để đảm bảo tính toàn vẹn của đồ thị tri thức."
    )
    
    # Bảng thống kê Node
    p_stats_node = doc.add_paragraph()
    p_stats_node.paragraph_format.space_before = Pt(8)
    p_stats_node.paragraph_format.space_after = Pt(4)
    p_stats_node.paragraph_format.keep_with_next = True
    r_title_n = p_stats_node.add_run("Bảng 1: Thống kê các loại Thực thể (Nodes) trong Neo4j")
    r_title_n.font.name = 'Times New Roman'
    r_title_n.bold = True
    r_title_n.font.size = Pt(14.5)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10.5)
    r_title_n.font.color.rgb = PRIMARY_COLOR
    
    nodes_data = [
        ("Loại Node (Label)", "Mô tả vai trò trong hệ thống", "Số lượng"),
        ("Law", "Văn bản Luật gốc (Dân sự, Hình sự, Lao động, Hôn nhân)", "4"),
        ("Chapter", "Các Chương trong các bộ luật", "79"),
        ("Article", "Các Điều luật cụ thể", "1.467"),
        ("Clause", "Các Khoản trong từng Điều luật", "984"),
        ("Concept", "Khái niệm, thuật ngữ pháp lý trích xuất bởi LLM", "5.906"),
        ("Actor", "Chủ thể chịu tác động hoặc thực hiện hành vi", "1.465"),
        ("Action", "Hành vi pháp lý, nghĩa vụ, quyền lợi", "5.347"),
        ("Chunk", "Đoạn văn bản nhỏ đã phân mảnh lưu Vector", "2.183"),
        ("Tổng số Nodes", "Toàn bộ thực thể đã lập chỉ mục", "17.435")
    ]
    
    table_n = doc.add_table(rows=len(nodes_data), cols=3)
    table_n.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_data in enumerate(nodes_data):
        row = table_n.rows[r_idx]
        is_header = (r_idx == 0)
        is_total = (r_idx == len(nodes_data) - 1)
        
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(3.4)
        row.cells[2].width = Inches(1.2)
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_before = Pt(4)
            p_cell.paragraph_format.space_after = Pt(4)
            if c_idx == 2:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p_cell.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10)
            
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "000000")  # Đầu bảng màu đen
            elif is_total:
                run.bold = True
                run.font.color.rgb = PRIMARY_COLOR
                set_cell_background(cell, "EAEAEA")
            else:
                run.font.color.rgb = TEXT_COLOR
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F9FBFD")
                    
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    
    # Bảng thống kê Relationship
    p_stats_rel = doc.add_paragraph()
    p_stats_rel.paragraph_format.space_before = Pt(8)
    p_stats_rel.paragraph_format.space_after = Pt(4)
    p_stats_rel.paragraph_format.keep_with_next = True
    r_title_r = p_stats_rel.add_run("Bảng 2: Thống kê các loại Quan hệ (Relationships) trong Neo4j")
    r_title_r.font.name = 'Times New Roman'
    r_title_r.bold = True
    r_title_r.font.size = Pt(14.5)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10.5)
    r_title_r.font.color.rgb = PRIMARY_COLOR
    
    rels_data = [
        ("Loại Quan hệ (Edge)", "Thực thể liên quan (Source -> Target)", "Số lượng"),
        ("HAS_CHAPTER", "Law -> Chapter (Chứa chương)", "79"),
        ("HAS_ARTICLE", "Chapter/Law -> Article (Chứa điều)", "1.467"),
        ("HAS_CLAUSE", "Article -> Clause (Chứa khoản)", "891"),
        ("IN_CHUNK", "Article/Clause -> Chunk (Nằm trong chunk)", "3.083"),
        ("REFERENCES", "Article -> Article (Dẫn chiếu trong cùng luật)", "1.556"),
        ("CROSS_LAW_REF", "Article -> Law (Dẫn chiếu chéo bộ luật khác)", "26"),
        ("DEFINES", "Article/Clause -> Concept (Định nghĩa khái niệm)", "284"),
        ("MENTIONS", "Chunk -> Concept (Đề cập khái niệm)", "5.913"),
        ("RELATED_TO", "Concept -> Concept (Khái niệm liên quan)", "7.919"),
        ("REGULATES", "Action -> Actor / Actor -> Action (Quy định hành vi)", "11.467"),
        ("PERFORMS", "Actor -> Action (Thực hiện hành vi)", "6.441"),
        ("Tổng số Relationships", "Toàn bộ liên kết cấu trúc và ngữ nghĩa", "39.126")
    ]
    
    table_r = doc.add_table(rows=len(rels_data), cols=3)
    table_r.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_data in enumerate(rels_data):
        row = table_r.rows[r_idx]
        is_header = (r_idx == 0)
        is_total = (r_idx == len(rels_data) - 1)
        
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(3.2)
        row.cells[2].width = Inches(1.2)
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_before = Pt(4)
            p_cell.paragraph_format.space_after = Pt(4)
            if c_idx == 2:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p_cell.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10)
            
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "000000")  # Đầu bảng màu đen
            elif is_total:
                run.bold = True
                run.font.color.rgb = PRIMARY_COLOR
                set_cell_background(cell, "EAEAEA")
            else:
                run.font.color.rgb = TEXT_COLOR
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F9FBFD")
                    
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
    # -------------------------------------------------------------
    # 1.2. Tích hợp Đồ thị vào Hệ thống RAG (Hybrid Fusion RAG)
    # -------------------------------------------------------------
    h1_2 = doc.add_paragraph()
    h1_2.paragraph_format.space_before = Pt(18)
    h1_2.paragraph_format.space_after = Pt(4)
    h1_2.paragraph_format.keep_with_next = True
    run_h1_2 = h1_2.add_run("1.2. Tích hợp Đồ thị vào Hệ thống RAG (Hybrid Fusion RAG)")
    run_h1_2.font.name = 'Times New Roman'
    run_h1_2.font.size = Pt(17)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 13)
    run_h1_2.font.bold = True
    run_h1_2.font.color.rgb = ACCENT_COLOR
    
    p = doc.add_paragraph(
        "Nhằm vượt qua các giới hạn của hệ thống RAG truyền thống (chỉ dựa vào khoảng cách không gian vector "
        "dễ dẫn đến việc mất ngữ cảnh cấu trúc hoặc trích xuất thiếu các điều khoản liên quan gián tiếp), "
        "chúng tôi đã triển khai kiến trúc Hybrid Fusion RAG kết hợp Vector Search và Graph Retrieval. "
        "Quy trình xử lý một câu truy vấn gồm các bước chính sau:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    steps = [
        ("Tìm kiếm Vector (Vector Search): ", "Truy vấn trên cơ sở dữ liệu ChromaDB sử dụng mô hình embedding BAAI/bge-m3 chạy trên GPU (CUDA) để lấy ra Top-K các chunk tài liệu có độ tương đồng ngữ nghĩa cao nhất với câu hỏi."),
        ("Truy vấn Đồ thị (Graph Retrieval): ", "LLM nhận dạng các thực thể pháp lý trong câu hỏi, sau đó thực hiện các truy vấn đồ thị Cypher đa bước (Multi-hop) để lấy ra các node Điều/Khoản lân cận, các định nghĩa khái niệm liên quan và các điều luật được dẫn chiếu từ kết quả tìm kiếm."),
        ("Kết hợp và Rerank (Reciprocal Rank Fusion - RRF): ", "Các tài liệu/chunk thu hồi từ cả hai kênh Vector và Đồ thị được chấm điểm và sắp xếp lại bằng thuật toán RRF để giữ lại các thông tin chất lượng nhất, hạn chế trùng lặp."),
        ("Tổng hợp câu trả lời (LLM Generation): ", "Ngữ cảnh sau khi rerank được đưa vào mô hình Gemini cùng với Prompt hệ thống tối ưu hóa cho luật Việt Nam để tạo ra câu trả lời cuối cùng, đảm bảo trích dẫn chính xác Điều, Khoản, và Tên bộ luật.")
    ]
    
    for bold_text, normal_text in steps:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        
        r_bold = bp.add_run(bold_text)
        r_bold.font.name = 'Times New Roman'
        r_bold.bold = True
        r_bold.font.color.rgb = TEXT_COLOR
        
        r_norm = bp.add_run(normal_text)
        r_norm.font.name = 'Times New Roman'
        r_norm.font.color.rgb = TEXT_COLOR
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(4)
    
    img_path = Path("diagram.png")
    if img_path.exists():
        p_img.add_run().add_picture(str(img_path), width=Inches(5.6))
        
        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_caption.paragraph_format.space_after = Pt(12)
        r_cap = p_caption.add_run("Hình 1: Sơ đồ kiến trúc hệ thống Hybrid RAG kết hợp Vector DB và Graph DB")
        r_cap.font.name = 'Times New Roman'
        r_cap.font.size = Pt(13.5)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 9.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(102, 102, 102)
    else:
        p_img.add_run("[Không tìm thấy file ảnh diagram.png để nhúng vào báo cáo]")
        p_img.runs[0].font.color.rgb = RGBColor(255, 0, 0)
 
    # -------------------------------------------------------------
    # 1.3. Tạo Tập dữ liệu Chưng cất (Dataset Distillation & Split)
    # -------------------------------------------------------------
    h1_3 = doc.add_paragraph()
    h1_3.paragraph_format.space_before = Pt(18)
    h1_3.paragraph_format.space_after = Pt(4)
    h1_3.paragraph_format.keep_with_next = True
    run_h1_3 = h1_3.add_run("1.3. Tạo Tập dữ liệu Chưng cất (Dataset Distillation & Split)")
    run_h1_3.font.name = 'Times New Roman'
    run_h1_3.font.size = Pt(17)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 13)
    run_h1_3.font.bold = True
    run_h1_3.font.color.rgb = ACCENT_COLOR
    
    p = doc.add_paragraph(
        "Nhằm xây dựng tập dữ liệu chất lượng cao phục vụ cho quá trình tinh chỉnh (Fine-tuning) mô hình "
        "ngôn ngữ lớn sau này, hệ thống đã thực hiện quy trình chưng cất dữ liệu (Dataset Distillation) "
        "tự động. Từ các đoạn văn bản luật gốc, hệ thống kết hợp LLM sinh ra các câu hỏi thực tế và câu trả lời tương ứng. "
        "Quy trình được kiểm soát chặt chẽ qua các bước:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    distill_points = [
        ("Sinh dữ liệu có ngữ cảnh: ", "Sử dụng LLM sinh câu hỏi và đáp án dựa trực tiếp trên nội dung của từng Điều/Khoản luật để đảm bảo câu trả lời luôn trung thực và chính xác tuyệt đối theo nguồn luật."),
        ("Cơ chế khử trùng lặp (Deduplication): ", "Trong quá trình sinh, các câu hỏi trùng lặp được phát hiện và loại bỏ tự động bằng cách băm MD5 chuỗi câu hỏi (đã chuẩn hóa về dạng chữ thường và loại bỏ khoảng trắng thừa). Nhờ đó, loại bỏ được các câu trùng lặp sinh ra từ các seed tương đồng."),
        ("Phân chia tập dữ liệu chuẩn hóa: ", "Tập dữ liệu chưng cất hoàn chỉnh gồm 9.546 cặp Q&A được phân chia ngẫu nhiên có kiểm soát (Random Seed = 42) thành 3 tập Train, Validation và Test theo tỷ lệ 80% / 10% / 10% để đảm bảo tính khách quan trong đánh giá mô hình.")
    ]
    for bold_text, normal_text in distill_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        
        r_bold = bp.add_run(bold_text)
        r_bold.font.name = 'Times New Roman'
        r_bold.bold = True
        r_bold.font.color.rgb = TEXT_COLOR
        
        r_norm = bp.add_run(normal_text)
        r_norm.font.name = 'Times New Roman'
        r_norm.font.color.rgb = TEXT_COLOR
        
    # Bảng phân chia dữ liệu
    p_stats_distill = doc.add_paragraph()
    p_stats_distill.paragraph_format.space_before = Pt(8)
    p_stats_distill.paragraph_format.space_after = Pt(4)
    p_stats_distill.paragraph_format.keep_with_next = True
    r_title_d = p_stats_distill.add_run("Bảng 3: Kết quả phân chia tập dữ liệu chưng cất (Dataset Splits)")
    r_title_d.font.name = 'Times New Roman'
    r_title_d.bold = True
    r_title_d.font.size = Pt(14.5)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10.5)
    r_title_d.font.color.rgb = PRIMARY_COLOR
    
    distill_data = [
        ("Tập dữ liệu (Split)", "Tỷ lệ phân chia", "Số lượng câu hỏi (Q&A)", "Kích thước tệp (JSONL)"),
        ("Tập huấn luyện (Train Set)", "80,0%", "7.636", "88,67 MB"),
        ("Tập kiểm thử song song (Val Set)", "10,0%", "954", "10,98 MB"),
        ("Tập kiểm thử độc lập (Test Set)", "10,0%", "956", "11,00 MB"),
        ("Tổng cộng", "100,0%", "9.546", "110,65 MB")
    ]
    
    table_d = doc.add_table(rows=len(distill_data), cols=4)
    table_d.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_data in enumerate(distill_data):
        row = table_d.rows[r_idx]
        is_header = (r_idx == 0)
        is_total = (r_idx == len(distill_data) - 1)
        
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(1.2)
        row.cells[2].width = Inches(1.6)
        row.cells[3].width = Inches(1.4)
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_before = Pt(4)
            p_cell.paragraph_format.space_after = Pt(4)
            if c_idx >= 2:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p_cell.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 10)
            
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "000000")  # Đầu bảng màu đen
            elif is_total:
                run.bold = True
                run.font.color.rgb = PRIMARY_COLOR
                set_cell_background(cell, "EAEAEA")
            else:
                run.font.color.rgb = TEXT_COLOR
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F9FBFD")
                    
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
 
    # -------------------------------------------------------------
    # 1.4. Một số ví dụ câu hỏi và câu trả lời trong dữ liệu chưng cất
    # -------------------------------------------------------------
    h1_4 = doc.add_paragraph()
    h1_4.paragraph_format.space_before = Pt(18)
    h1_4.paragraph_format.space_after = Pt(4)
    h1_4.paragraph_format.keep_with_next = True
    run_h1_4 = h1_4.add_run("1.4. Một số ví dụ mẫu câu hỏi và câu trả lời chưng cất")
    run_h1_4.font.name = 'Times New Roman'
    run_h1_4.font.size = Pt(17)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 13)
    run_h1_4.font.bold = True
    run_h1_4.font.color.rgb = ACCENT_COLOR
 
    p_sample_intro = doc.add_paragraph(
        "Dưới đây là một số ví dụ câu hỏi thực tế được sinh bởi hệ thống kèm câu trả lời tham chiếu nguồn "
        "luật tương ứng từ tập dữ liệu chưng cất (Dataset Distillation) đã sinh thành công:"
    )
    p_sample_intro.paragraph_format.space_after = Pt(6)
 
    # Ví dụ 1
    p_ex1 = doc.add_paragraph()
    p_ex1.paragraph_format.space_before = Pt(6)
    p_ex1.paragraph_format.space_after = Pt(3)
    r_ex1_t = p_ex1.add_run("Ví dụ 1 (Lĩnh vực Lao động):")
    r_ex1_t.font.name = 'Times New Roman'
    r_ex1_t.font.bold = True
    r_ex1_t.font.size = Pt(15)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 11)
    
    p_ex1_q = doc.add_paragraph()
    p_ex1_q.paragraph_format.left_indent = Inches(0.2)
    p_ex1_q.paragraph_format.space_after = Pt(3)
    r_q1_lbl = p_ex1_q.add_run("- Câu hỏi: ")
    r_q1_lbl.font.name = 'Times New Roman'
    r_q1_lbl.font.bold = True
    r_q1_val = p_ex1_q.add_run(
        "Theo Bộ luật Lao động 2019, doanh nghiệp có quyền tạm đình chỉ công việc của người lao động trong trường hợp nào "
        "và có cần tham khảo ý kiến của tổ chức nào trước khi thực hiện không?"
    )
    r_q1_val.font.name = 'Times New Roman'
 
    p_ex1_a = doc.add_paragraph()
    p_ex1_a.paragraph_format.left_indent = Inches(0.2)
    p_ex1_a.paragraph_format.space_after = Pt(8)
    r_a1_lbl = p_ex1_a.add_run("- Trả lời: ")
    r_a1_lbl.font.name = 'Times New Roman'
    r_a1_lbl.font.bold = True
    r_a1_val = p_ex1_a.add_run(
        "Theo Khoản 1 Điều 128 Bộ luật Lao động 2019, doanh nghiệp (người sử dụng lao động) có quyền tạm đình chỉ "
        "công việc của người lao động khi vụ việc vi phạm có những tình tiết phức tạp nếu xét thấy để người lao động tiếp "
        "tục làm việc sẽ gây khó khăn cho việc xác minh. Trước khi thực hiện, doanh nghiệp bắt buộc phải tham khảo ý kiến "
        "của tổ chức đại diện người lao động tại cơ sở mà người lao động đang bị xem xét tạm đình chỉ công việc là thành viên."
    )
    r_a1_val.font.name = 'Times New Roman'
 
    # Ví dụ 2
    p_ex2 = doc.add_paragraph()
    p_ex2.paragraph_format.space_before = Pt(6)
    p_ex2.paragraph_format.space_after = Pt(3)
    r_ex2_t = p_ex2.add_run("Ví dụ 2 (Lĩnh vực Dân sự):")
    r_ex2_t.font.name = 'Times New Roman'
    r_ex2_t.font.bold = True
    r_ex2_t.font.size = Pt(15)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 11)
    
    p_ex2_q = doc.add_paragraph()
    p_ex2_q.paragraph_format.left_indent = Inches(0.2)
    p_ex2_q.paragraph_format.space_after = Pt(3)
    r_q2_lbl = p_ex2_q.add_run("- Câu hỏi: ")
    r_q2_lbl.font.name = 'Times New Roman'
    r_q2_lbl.font.bold = True
    r_q2_val = p_ex2_q.add_run(
        "Anh A và chị B có con chung là C (10 tuổi). Sau đó, anh A muốn thay đổi họ cho con từ họ của mình sang họ của chị B. "
        "Việc thay đổi họ này có được phép không và cần có những điều kiện gì?"
    )
    r_q2_val.font.name = 'Times New Roman'
 
    p_ex2_a = doc.add_paragraph()
    p_ex2_a.paragraph_format.left_indent = Inches(0.2)
    p_ex2_a.paragraph_format.space_after = Pt(8)
    r_a2_lbl = p_ex2_a.add_run("- Trả lời: ")
    r_a2_lbl.font.name = 'Times New Roman'
    r_a2_lbl.font.bold = True
    r_a2_val = p_ex2_a.add_run(
        "Việc thay đổi họ từ họ cha đẻ sang họ mẹ đẻ là được phép theo quy định tại Điều 27 Khoản 1 Điểm a Luật Dân sự 2015. "
        "Về điều kiện đi kèm, do cháu C đã 10 tuổi (thuộc nhóm từ đủ chín tuổi trở lên), căn cứ theo Điều 27 Khoản 2 Luật Dân sự 2015, "
        "việc thay đổi họ phải có sự đồng ý của bản thân cháu C."
    )
    r_a2_val.font.name = 'Times New Roman'
            
    # -------------------------------------------------------------
    # 2. KẾ HOẠCH THỰC HIỆN TIẾP THEO
    # -------------------------------------------------------------
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    run_h2 = h2.add_run("2. Kế hoạch thực hiện tiếp theo")
    run_h2.font.name = 'Times New Roman'
    run_h2.font.size = Pt(20)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph("Trong giai đoạn tiếp theo, dự án tập trung hoàn thiện các nội dung sau:")
    p.paragraph_format.space_after = Pt(6)
    
    plans = [
        ("2.1. Huấn luyện tinh chỉnh mô hình ngôn ngữ (Fine-tuning LLM / Train Model)",
         ["Sử dụng tập dữ liệu huấn luyện chưng cất (Train Set - 7.636 mẫu Q&A) để tiến hành tinh chỉnh các mô hình ngôn ngữ lớn chuyên sâu (như Qwen-2.5-7B, LLaMA-3-8B hoặc PhoGPT) chuyên sâu về pháp luật Việt Nam.",
          "Áp dụng kỹ thuật huấn luyện hiệu quả tham số QLoRA / LoRA nhằm tối ưu hóa bộ nhớ GPU và thời gian huấn luyện.",
          "Theo dõi sát chỉ số training loss và validation loss (trên tập Val Set gồm 954 mẫu) để phát hiện sớm và hạn chế hiện tượng overfitting."]),
        ("2.2. Xây dựng giao diện người dùng (Frontend - FE) và Tích hợp hệ thống",
         ["Thiết kế và dựng cấu trúc dự án Frontend (FE) hoàn chỉnh sử dụng framework React/Vite hoặc Gradio Web App tùy chỉnh.",
          "Xây dựng các màn hình giao diện tương tác: Màn hình Đăng nhập/Đăng ký cho người dùng, Giao diện trò chuyện trực quan (Chatbot UI) hỗ trợ hiển thị lịch sử chat và xuất nguồn dẫn chiếu, và Giao diện Quản lý tài liệu cho phép Admin upload/cập nhật văn bản luật và hiển thị các số liệu thống kê đồ thị thời gian thực.",
          "Đồng bộ hóa toàn bộ các luồng (flow) dữ liệu giữa Client Frontend và các REST API được viết bằng FastAPI ở Backend."]),
        ("2.3. Tối ưu hóa RAG và Truy vấn Đồ thị",
         ["Thực hiện tinh chỉnh các siêu tham số retrieval: thử nghiệm các giá trị Top-K khác nhau và ngưỡng tương đồng để đạt độ bao phủ ngữ cảnh tối ưu.",
          "Cải tiến thuật toán Hybrid Fusion (RRF) thông qua việc điều chỉnh các tham số fusion để cân bằng sự đóng góp từ Neo4j Graph DB và Chroma Vector DB.",
          "Tối ưu hóa tốc độ các câu truy vấn Cypher đa bước trên Neo4j bằng cách tạo index trên các thuộc tính của nút (Node properties) và cấu trúc lại quan hệ."]),
        ("2.4. Đánh giá chất lượng toàn diện",
         ["Đánh giá tự động trên tập Test (956 mẫu) sử dụng framework RAGAS đối với các chỉ số quan trọng: Độ chính xác ngữ cảnh (Context Precision), Độ phủ ngữ cảnh (Context Recall), Độ trung thực câu trả lời (Faithfulness) và Độ liên quan của câu trả lời (Answer Relevance).",
          "Đo lường các chỉ số sinh văn bản chuẩn bao gồm BLEU và ROUGE so sánh với đáp án chuẩn trong bộ dữ liệu.",
          "Tổ chức kiểm thử mù (Blind Test) với sự tham gia của các chuyên gia pháp lý để đánh giá độ tin cậy và tính ứng dụng thực tiễn của hệ thống RAG trong các tình huống thực tế."])
    ]
    
    for sub_title, steps_list in plans:
        h2_sub = doc.add_paragraph()
        h2_sub.paragraph_format.space_before = Pt(10)
        h2_sub.paragraph_format.space_after = Pt(4)
        h2_sub.paragraph_format.keep_with_next = True
        r_sub = h2_sub.add_run(sub_title)
        r_sub.font.name = 'Times New Roman'
        r_sub.font.size = Pt(16)  # Tăng thêm 2 pt (tổng cộng tăng 4 pt từ 12)
        r_sub.font.bold = True
        r_sub.font.color.rgb = ACCENT_COLOR
        
        for step in steps_list:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(3)
            r_st = bp.add_run(step)
            r_st.font.name = 'Times New Roman'
            r_st.font.color.rgb = TEXT_COLOR
            
    # Lưu tài liệu
    doc_name = "Bao_cao_Tien_do.docx"
    try:
        doc.save(doc_name)
        print(f"Báo cáo tiến độ đã được cập nhật thành công (Times New Roman, cỡ chữ tăng 2, màu đen) tại: {doc_name}")
    except PermissionError:
        backup_name = "Bao_cao_Tien_do_v2.docx"
        doc.save(backup_name)
        print(f"Không thể ghi đè '{doc_name}' do tệp đang mở ở chương trình khác.")
        print(f"Báo cáo đã được lưu tạm thời tại: {backup_name}")

if __name__ == "__main__":
    main()
