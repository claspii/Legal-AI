"""
Script tự động sinh tệp Word (.docx) báo cáo tiến độ đồ án tốt nghiệp chi tiết,
được thiết kế theo tiêu chuẩn định dạng học thuật Việt Nam (Times New Roman, lề chuẩn, bảng biểu có padding và màu nền tinh tế).
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

BASE_DIR = Path(__file__).parent.parent

def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt padding cho ô trong bảng (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_callout(doc, text):
    """Tạo callout box nổi bật có viền trái đen dày và nền xám nhạt."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="12" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F8"/>')
    pPr.append(shd)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_placeholder_box(doc, title_text, description):
    """Tạo một khung placeholder viền đứt nét để người dùng chèn ảnh chụp màn hình sau."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.0)
    
    # Định dạng viền đứt nét màu Slate xám cho khung ảnh
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="dashed" w:sz="8" w:space="0" w:color="64748B"/>'
        f'<w:left w:val="dashed" w:sz="8" w:space="0" w:color="64748B"/>'
        f'<w:bottom w:val="dashed" w:sz="8" w:space="0" w:color="64748B"/>'
        f'<w:right w:val="dashed" w:sz="8" w:space="0" w:color="64748B"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    set_cell_background(cell, "F8FAFC") # Nền Slate nhạt cực sang
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_icon = p.add_run("🖼️ [HÌNH ẢNH MINH HỌA - CHÈN TẠI ĐÂY]\n")
    run_icon.bold = True
    run_icon.font.size = Pt(12)
    run_icon.font.color.rgb = RGBColor(14, 165, 233) # Teal
    
    run_title = p.add_run(f"{title_text}\n")
    run_title.bold = True
    run_title.font.size = Pt(13)
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    run_desc = p.add_run(description)
    run_desc.font.size = Pt(11)
    run_desc.font.italic = True
    run_desc.font.color.rgb = RGBColor(100, 116, 139)
    
    # Thêm caption cho ảnh ở dưới bảng
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"Hình: {title_text}")
    run_cap.font.name = 'Times New Roman'
    run_cap.font.size = Pt(11)
    run_cap.font.italic = True
    run_cap.font.color.rgb = RGBColor(100, 116, 139)

def main():
    doc = docx.Document()
    
    # L lề trang tiêu chuẩn (1 inch = 2.54 cm cho cả 4 phía)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Màu sắc chủ đạo học thuật (Màu đen tiêu chuẩn)
    PRIMARY_COLOR = RGBColor(0, 0, 0)
    TEXT_COLOR = RGBColor(0, 0, 0)
    ACCENT_COLOR = RGBColor(0, 0, 0)
    
    # Cấu hình Default Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(13)  # Cỡ chữ 13 chuẩn tiểu luận
    style_normal.font.color.rgb = TEXT_COLOR
    
    # -------------------------------------------------------------
    # TIÊU ĐỀ BÁO CÁO
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(6)
    
    title_run = title_p.add_run("BÁO CÁO TIẾN ĐỘ ĐỒ ÁN TỐT NGHIỆP CHI TIẾT")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle_p.add_run("Phát triển Hệ thống Hỏi đáp Pháp luật Hybrid RAG + Knowledge Graph & Tinh chỉnh Mô hình (SFT)")
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = ACCENT_COLOR
    
    # -------------------------------------------------------------
    # 1. TỔNG QUAN VÀ MỤC TIÊU ĐỀ TÀI
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True
    run_h1 = h1.add_run("1. Tổng quan và Mục tiêu đề tài")
    run_h1.font.name = 'Times New Roman'
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "Đề tài hướng tới việc xây dựng một hệ thống Trợ lý ảo tư vấn Pháp luật Việt Nam có độ tin cậy cực cao, "
        "khắc phục triệt để hiện tượng ảo tưởng (hallucination) của các mô hình ngôn ngữ lớn (LLM). Hệ thống tích hợp "
        "kỹ thuật truy hồi tăng cường đồ thị (Graph RAG) kết hợp tìm kiếm vector, đồng thời tinh chỉnh một mô hình "
        "chuyên biệt (Domain-specific SFT LLM) chuyên sâu về văn bản pháp lý Việt Nam, phục vụ người dân và các luật sư tra cứu thông tin nhanh chóng."
    )
    p.paragraph_format.space_after = Pt(6)
    
    # -------------------------------------------------------------
    # 2. CÁC NỘI DUNG CHÍNH ĐÃ HOÀN THÀNH (CHI TIẾT)
    # -------------------------------------------------------------
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    run_h2 = h2.add_run("2. Các nội dung chính đã hoàn thành")
    run_h2.font.name = 'Times New Roman'
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = PRIMARY_COLOR
    
    # 2.1. Phân mảnh & Lập chỉ mục Vector (ChromaDB)
    h2_1 = doc.add_paragraph()
    h2_1.paragraph_format.space_before = Pt(12)
    h2_1.paragraph_format.space_after = Pt(4)
    h2_1.paragraph_format.keep_with_next = True
    run_h2_1 = h2_1.add_run("2.1. Xây dựng bộ tìm kiếm Vector ngữ nghĩa (Vector Store)")
    run_h2_1.font.bold = True
    run_h2_1.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Chúng tôi đã phát triển bộ nạp tài liệu tự động trích xuất cấu trúc văn bản luật từ 4 bộ luật cốt lõi "
        "(Dân sự, Hình sự, Hôn nhân và Gia đình, Lao động). Hệ thống áp dụng chiến lược phân mảnh lai (Hybrid Chunking Strategy): "
        "các điều luật ngắn được giữ nguyên làm một chunk để bảo toàn ngữ cảnh hoàn chỉnh, các điều luật dài có cấu trúc phức tạp "
        "được phân mảnh mịn hơn theo từng Khoản. Việc nhúng từ (Embedding) sử dụng mô hình BAAI/bge-m3 hiệu năng cao trên GPU (CUDA) "
        "và lập chỉ mục trong cơ sở dữ liệu vector ChromaDB."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # 2.2. Cơ sở dữ liệu đồ thị tri thức (Neo4j Knowledge Graph)
    h2_2 = doc.add_paragraph()
    h2_2.paragraph_format.space_before = Pt(12)
    h2_2.paragraph_format.space_after = Pt(4)
    h2_2.paragraph_format.keep_with_next = True
    run_h2_2 = h2_2.add_run("2.2. Thiết lập cơ sở dữ liệu đồ thị tri thức (Neo4j Graph)")
    run_h2_2.font.bold = True
    run_h2_2.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Một đồ thị tri thức Neo4j được thiết kế và lập chỉ mục chặt chẽ để ánh xạ cấu trúc logic của văn bản luật: "
        "Nút cha Luật chứa các Chương, Chương chứa các Điều, Điều chứa các Khoản và Chunk. Đồng thời áp dụng "
        "biểu thức chính quy (Regex) và gọi LLM (Gemini API có cache) để khai thác các cạnh dẫn chiếu (REFERENCES) chéo giữa các điều, "
        "và các quan hệ ngữ nghĩa tinh thể (Định nghĩa khái niệm Concept, Chủ thể Actor, Hành vi Action)."
    )
    p.paragraph_format.space_after = Pt(6)
    
    add_callout(
        doc,
        "Đồ thị tri thức hiện đã được lập chỉ mục hoàn chỉnh với hơn 17.400 Thực thể (Nodes) và hơn 39.100 Quan hệ (Relationships), "
        "đảm bảo thời gian truy vấn đa bước (Multi-hop Query) cực kỳ tối ưu."
    )
    
    # Chèn bảng thống kê thực thể đồ thị
    table_n = doc.add_table(rows=5, cols=3)
    table_n.alignment = WD_TABLE_ALIGNMENT.CENTER
    nodes_stats = [
        ("Thành phần thực thể", "Mô tả vai trò trong hệ thống", "Số lượng"),
        ("Nút cấu trúc (Law, Chapter, Article, Clause)", "Ánh xạ sơ đồ cấu trúc chương điều khoản luật thô", "2.534"),
        ("Nút ngữ nghĩa (Concept, Actor, Action)", "Thực thể định nghĩa, chủ thể pháp luật và hành vi pháp lý", "12.718"),
        ("Nút liên kết dữ liệu (Chunk)", "Các đoạn văn bản đã băm nhúng lưu trữ vector", "2.183"),
        ("Tổng số thực thể (Total Nodes)", "Quy mô cơ sở dữ liệu tri thức tích hợp", "17.435")
    ]
    for r_idx, r_data in enumerate(nodes_stats):
        row = table_n.rows[r_idx]
        is_header = (r_idx == 0)
        is_total = (r_idx == len(nodes_stats) - 1)
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(2.8)
        row.cells[2].width = Inches(1.0)
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_before = Pt(4)
            p_cell.paragraph_format.space_after = Pt(4)
            if c_idx == 2:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p_cell.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "000000")
            elif is_total:
                run.bold = True
                set_cell_background(cell, "F1F5F9")
            else:
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    add_placeholder_box(
        doc,
        "Đồ thị trực quan hóa cấu trúc thực thể trong Neo4j",
        "Hãy chụp lại sơ đồ Neo4j Browser hiển thị các nhãn thực thể và quan hệ (đã liên kết cấu trúc và ngữ nghĩa) và dán vào đây."
    )
    
    # 2.3. Hợp nhất truy hồi Hybrid RAG (RRF)
    h2_3 = doc.add_paragraph()
    h2_3.paragraph_format.space_before = Pt(12)
    h2_3.paragraph_format.space_after = Pt(4)
    h2_3.paragraph_format.keep_with_next = True
    run_h2_3 = h2_3.add_run("2.3. Giải thuật hợp nhất truy hồi Hybrid RAG")
    run_h2_3.font.bold = True
    run_h2_3.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Nhóm đã xây dựng mô-đun phối hợp truy xuất dữ liệu từ hai kênh song song. Kênh 1 sử dụng ChromaDB để tìm "
        "các đoạn tương đồng cosine cao nhất. Kênh 2 quét đồ thị Neo4j dựa trên thực thể trích từ câu hỏi và mở rộng "
        "các điều dẫn chiếu lân cận từ hạt giống tìm được. Sau đó, thuật toán Reciprocal Rank Fusion (RRF) được sử dụng "
        "để tính điểm tổng hợp và xếp hạng lại toàn bộ các chunk tài liệu, gom lại thành một ngữ cảnh đầu vào tối ưu nhất cho LLM."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_placeholder_box(
        doc,
        "Sơ đồ kiến trúc giải thuật Hybrid RAG kết hợp Vector + Graph",
        "Dán hình ảnh sơ đồ khối/kiến trúc luồng xử lý câu truy vấn từ câu hỏi người dùng qua hai kênh và hợp nhất bằng RRF."
    )
    
    # 2.4. Chưng cất & Tinh chỉnh dữ liệu
    h2_4 = doc.add_paragraph()
    h2_4.paragraph_format.space_before = Pt(12)
    h2_4.paragraph_format.space_after = Pt(4)
    h2_4.paragraph_format.keep_with_next = True
    run_h2_4 = h2_4.add_run("2.4. Chưng cất tự động và phân chia tập dữ liệu huấn luyện")
    run_h2_4.font.bold = True
    run_h2_4.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Hệ thống đã triển khai quy trình chưng cất tự động (Dataset Distillation). Đầu tiên, sinh câu hỏi đa dạng góc độ "
        "dựa theo từng khoản luật (seed). Sau đó, dùng RAG truy hồi ngữ cảnh chính xác nhất và gọi LLM thông minh sinh "
        "chuỗi suy luận từng bước (Reasoning) cùng đáp án chuẩn có trích dẫn điều khoản luật. Tệp dữ liệu chưng cất thô gồm "
        "hơn 9.500 mẫu Q&A đã được băm MD5 để loại bỏ hoàn toàn các câu hỏi trùng lặp, sau đó được chia ngẫu nhiên thành "
        "3 tập: Train (80%), Validation (10%), và Test (10%)."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Bảng chia dữ liệu
    table_d = doc.add_table(rows=4, cols=3)
    table_d.alignment = WD_TABLE_ALIGNMENT.CENTER
    distill_stats = [
        ("Tập dữ liệu (Split)", "Tỷ lệ phân chia", "Số lượng câu hỏi (Q&A)"),
        ("Tập huấn luyện (Train Set)", "80,0%", "7.636"),
        ("Tập đánh giá trung gian (Val Set)", "10,0%", "954"),
        ("Tập đánh giá độc lập (Test Set)", "10,0%", "956")
    ]
    for r_idx, r_data in enumerate(distill_stats):
        row = table_d.rows[r_idx]
        is_header = (r_idx == 0)
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(2.0)
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_before = Pt(4)
            p_cell.paragraph_format.space_after = Pt(4)
            if c_idx >= 1:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p_cell.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "000000")
            else:
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    
    # 2.5. Tinh chỉnh mô hình (SFT) & Đồ thị Loss
    h2_5 = doc.add_paragraph()
    h2_5.paragraph_format.space_before = Pt(12)
    h2_5.paragraph_format.space_after = Pt(4)
    h2_5.paragraph_format.keep_with_next = True
    run_h2_5 = h2_5.add_run("2.5. Huấn luyện tinh chỉnh mô hình SFT & Kết quả hội tụ")
    run_h2_5.font.bold = True
    run_h2_5.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Sử dụng thư viện Unsloth tối ưu hóa, nhóm đã thực hiện tinh chỉnh (Fine-tuning) mô hình ngôn ngữ Qwen-3.5 bằng kỹ thuật "
        "QLoRA / LoRA hiệu năng cao trên tập Train gồm 7.636 mẫu Q&A. Tệp lịch sử huấn luyện trainer_state.json đã được phân tích "
        "và trực quan hóa thành công. Kết quả chỉ ra rằng mô hình đạt độ hội tụ cực tốt và đạt giá trị mất mát trên tập đánh giá "
        "(Validation Loss) thấp nhất là 0.2858 tại Step 1800. Việc dừng sớm và khôi phục checkpoint tốt nhất này giúp mô hình "
        "đạt trạng thái hội tụ tối ưu mà không bị overfitting."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_placeholder_box(
        doc,
        "Đồ thị Phân tích tiến trình Huấn luyện LoRA (Loss Curves, LR, Grad Norm)",
        "Hãy dán hình ảnh dashboard 4 biểu đồ Slate/Neon được xuất tự động từ scripts/plot_training_results.py vào đây."
    )
    
    # 2.6. Triển khai Cloud API
    h2_6 = doc.add_paragraph()
    h2_6.paragraph_format.space_before = Pt(12)
    h2_6.paragraph_format.space_after = Pt(4)
    h2_6.paragraph_format.keep_with_next = True
    run_h2_6 = h2_6.add_run("2.6. Đóng gói mô hình GGUF và triển khai Cloud API")
    run_h2_6.font.bold = True
    run_h2_6.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Mô hình tinh chỉnh lora tại checkpoint-1800 tốt nhất được merge trực tiếp vào base model và export ra định dạng GGUF Q8_0 "
        "để tối ưu hóa bộ nhớ và tốc độ suy luận. GGUF model đã được tải lên kho lưu trữ Hugging Face và triển khai "
        "thành công trên máy chủ đám mây công khai (Cloud VM). Chúng tôi sử dụng ngrok để thiết lập đường truyền public URL bảo mật "
        "đầu ra dạng OpenAI-compatible endpoint phục vụ tích hợp trực tiếp vào ứng dụng chính."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_placeholder_box(
        doc,
        "Giao diện lưu trữ mô hình trên Hugging Face & Console khởi chạy máy chủ Cloud",
        "Hãy chụp ảnh màn hình Hugging Face Repository chứa tệp model GGUF Q8_0 và hình ảnh terminal/console chạy vLLM/llama.cpp server trên cloud rồi dán vào đây."
    )
    
    # 2.7. Giao diện Gradio và Streaming Reasoning
    h2_7 = doc.add_paragraph()
    h2_7.paragraph_format.space_before = Pt(12)
    h2_7.paragraph_format.space_after = Pt(4)
    h2_7.paragraph_format.keep_with_next = True
    run_h2_7 = h2_7.add_run("2.7. Tích hợp giao diện Chatbot nâng cao và Streaming Reasoning")
    run_h2_7.font.bold = True
    run_h2_7.font.size = Pt(14)
    
    p = doc.add_paragraph(
        "Chúng tôi đã phát triển và hoàn thiện giao diện Chatbot tương tác thời gian thực trên nền tảng Gradio 6.0 mới nhất. "
        "Giao diện tích hợp đầy đủ bảng điều khiển tham số (Temperature, Max tokens, Top-p, Top-k, và các Penalty) để "
        "điều khiển API tinh chỉnh linh hoạt. Hệ thống hỗ trợ chế độ Streaming Answer (câu trả lời sinh ra chữ nào hiện chữ đó) "
        "và đặc biệt là Streaming Reasoning: quá trình suy luận logic từng bước của mô hình chuyên sâu được gom lại trong một "
        "dropdown '<details>' co giãn, hiển thị động 🧠 Đang suy luận (Reasoning)... và tự thu gọn khi hoàn thành để tạo trải nghiệm trực quan cao cấp."
    )
    p.paragraph_format.space_after = Pt(8)
    
    add_placeholder_box(
        doc,
        "Giao diện Chatbot tương tác thực tế với luồng Streaming Reasoning",
        "Chụp ảnh màn hình giao diện Gradio khi đang chạy thực tế, hiển thị chi tiết hộp dropdown suy luận đang chạy và kết quả phản hồi trích dẫn điều khoản luật chi tiết."
    )

    # -------------------------------------------------------------
    # 3. KẾ HOẠCH THỰC HIỆN TIẾP THEO
    # -------------------------------------------------------------
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(20)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True
    run_h3 = h3.add_run("3. Kế hoạch thực hiện tiếp theo")
    run_h3.font.name = 'Times New Roman'
    run_h3.font.size = Pt(16)
    run_h3.font.bold = True
    run_h3.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph(
        "Để chuẩn bị tốt nhất cho giai đoạn bảo vệ đồ án tốt nghiệp, nhóm sẽ tập trung vào các nội dung trọng tâm sau:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    plans = [
        ("Tối ưu hóa các siêu tham số truy xuất Hybrid RAG: ", "Thử nghiệm thêm các trọng số fusion khác nhau giữa kênh Vector và Graph để tăng cường chất lượng nguồn dẫn chứng."),
        ("Đánh giá chất lượng tự động: ", "Áp dụng framework RAGAS để đo lường độ trung thực (Faithfulness) và độ liên quan của phản hồi. Sử dụng các chỉ số BLEU và ROUGE để so sánh với đáp án chuẩn trong tập dữ liệu Test."),
        ("Tối ưu tốc độ phản hồi đồ thị: ", "Lập thêm các chỉ số phụ trên các thuộc tính của các Concept và Actor trong Neo4j để đẩy nhanh tốc độ truy vấn Cypher đa bước."),
        ("Hoàn thiện giao diện và Đóng gói: ", "Đóng gói ứng dụng chính bằng Docker để dễ dàng triển khai phục vụ người dùng thử nghiệm thực tế.")
    ]
    for bold_text, normal_text in plans:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r_bold = bp.add_run(bold_text)
        r_bold.font.name = 'Times New Roman'
        r_bold.bold = True
        r_bold.font.color.rgb = TEXT_COLOR
        r_norm = bp.add_run(normal_text)
        r_norm.font.name = 'Times New Roman'
        r_norm.font.color.rgb = TEXT_COLOR

    # Lưu tệp báo cáo tiến độ chi tiết
    doc_path = BASE_DIR / "docs" / "Bao_cao_Tien_do_Datn_Full.docx"
    doc.save(str(doc_path))
    print(f"Đã tạo báo cáo tiến độ chi tiết thành công tại: {doc_path}")

if __name__ == "__main__":
    main()
