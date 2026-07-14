"""
Document generator utility: converts markdown text to DOCX and PDF formats.
Handles programmatic installation of required packages.
"""

import sys
import os
import tempfile
import subprocess
from io import BytesIO
from loguru import logger

# Try importing dependencies, install if missing
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    logger.info("Installing python-docx programmatically...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        logger.error(f"Failed to install python-docx: {e}")




def markdown_to_docx(md_text: str, style_guide: dict = None) -> bytes:
    """
    Convert markdown legal text into a styled DOCX document.
    """
    style_guide = style_guide or {}
    font_name = style_guide.get("font_name", "Times New Roman")
    
    # Body font size: standard Vietnamese legal documents are 13pt or 14pt
    try:
        size_str = style_guide.get("font_size_body", "13").replace("pt", "")
        font_size_pt = int(size_str)
    except ValueError:
        font_size_pt = 13

    doc = Document()
    
    # Set page margins conforming to Nghị định 30/2020/NĐ-CP: 
    # Top/Bottom: 2cm, Left: 3cm (for binding), Right: 1.5cm - 2cm (we use 1.5cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)      # 2.0 cm
        section.bottom_margin = Inches(0.79)   # 2.0 cm
        section.left_margin = Inches(1.18)     # 3.0 cm (for binding)
        section.right_margin = Inches(0.59)    # 1.5 cm

    # Set base style
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)

    lines = md_text.split('\n')
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()
        
        # Table parsing
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            continue
        elif in_table:
            # End of table, process table rows
            _add_docx_table(doc, table_rows, font_name, font_size_pt)
            table_rows = []
            in_table = False

        if not stripped:
            continue

        # Horizontal Rule (Divider)
        if stripped == '---' or stripped == '***' or stripped == '___' or (stripped.startswith('-') and all(c == '-' for c in stripped) and len(stripped) >= 3):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("—" * 15)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            # Make the divider gray
            run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            continue

        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt + 4)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            # If it's a section, center it
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if stripped[3:].isupper() else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt + 2)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt + 1)
        # Bullet lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            _add_formatted_text(p, stripped[2:], font_name, font_size_pt)
        # Numbered lists
        elif stripped.startswith('1. ') or stripped.startswith('2. ') or stripped.startswith('3. '):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            dot_idx = stripped.find('. ')
            _add_formatted_text(p, stripped[dot_idx+2:], font_name, font_size_pt)
        # Blockquotes
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            _add_formatted_text(p, stripped[2:], font_name, font_size_pt, italic=True)
        # Normal paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            
            # Alignments heuristic for Vietnamese Legal Documents
            if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(2)
                _add_formatted_text(p, stripped, font_name, font_size_pt, bold_all=True)
            elif "Độc lập - Tự do - Hạnh phúc" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                _add_formatted_text(p, stripped, font_name, font_size_pt, bold_all=True)
                
                # Add standard national slogan underline
                p_line = doc.add_paragraph()
                p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_line.paragraph_format.space_before = Pt(0)
                p_line.paragraph_format.space_after = Pt(12)
                run_line = p_line.add_run("—" * 8)
                run_line.bold = True
                run_line.font.name = font_name
                run_line.font.size = Pt(font_size_pt)
            elif "ĐƠN KHỞI KIỆN" in stripped or "HỢP ĐỒNG" in stripped or "GIẤY ỦY QUYỀN" in stripped or "THỎA THUẬN" in stripped or "BIÊN BẢN" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                _add_formatted_text(p, stripped, font_name, font_size_pt + 2, bold_all=True)
            elif stripped.startswith('|') and stripped.endswith('|'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_formatted_text(p, stripped, font_name, font_size_pt)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Heuristic: apply first line indentation for long normal body paragraphs
                if len(stripped) > 80:
                    p.paragraph_format.first_line_indent = Inches(0.49) # ~1.25 cm
                _add_formatted_text(p, stripped, font_name, font_size_pt)

    # Process remaining table if exists
    if in_table and table_rows:
        _add_docx_table(doc, table_rows, font_name, font_size_pt)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


def _add_formatted_text(paragraph, text, font_name, base_size, italic=False, bold_all=False):
    """Parses simple markdown formatting (**bold**, *italic*) and adds to paragraph."""
    import re
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            if bold_all:
                run.bold = True
        else:
            run = paragraph.add_run(token)
            if bold_all:
                run.bold = True
            
        run.font.name = font_name
        run.font.size = Pt(base_size)
        if italic:
            run.italic = True


def _add_docx_table(doc, rows, font_name, base_size):
    """Parses markdown table lines and adds an native Word Table."""
    parsed_rows = []
    for r in rows:
        cells = [c.strip() for c in r.split('|')]
        # If line starts and ends with |, cells[0] and cells[-1] will be empty
        if len(cells) > 1 and cells[0] == '':
            cells = cells[1:]
        if len(cells) > 0 and cells[-1] == '':
            cells = cells[:-1]
        
        # Skip divider rows (e.g. |:---:|:---:|)
        if cells and all(c.replace(':', '').replace('-', '').strip() == '' for c in cells):
            continue
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    # Check if signature table or borderless block (contains key sign keywords)
    is_signature_table = False
    for row_cells in parsed_rows:
        for cell_val in row_cells:
            val_upper = cell_val.upper()
            if any(kw in val_upper for kw in ["BÊN A", "BÊN B", "ĐẠI DIỆN", "NGƯỜI ỦY QUYỀN", "NGƯỜI ĐƯỢC ỦY QUYỀN", "KÝ TÊN", "KÝ VÀ GHI RÕ"]):
                is_signature_table = True
                break
        if is_signature_table:
            break

    num_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    
    if is_signature_table:
        try:
            table.style = 'Normal Table'
        except KeyError:
            pass
    else:
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass

    # Center align table
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_cells in enumerate(parsed_rows):
        row = table.rows[r_idx]
        for c_idx, cell_val in enumerate(row_cells):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                
                if is_signature_table:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Make headers or main terms bold in signature block
                    _add_formatted_text(p, cell_val, font_name, base_size, bold_all=(r_idx == 0))
                else:
                    # Normal tables: make header row bold
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _add_formatted_text(p, cell_val, font_name, base_size, bold_all=(r_idx == 0))



